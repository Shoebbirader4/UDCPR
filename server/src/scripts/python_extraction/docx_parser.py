"""
DOCX Parser - Extract structured content from UDCPR/Mumbai-DCPR DOCX files
"""
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import re
from typing import List, Dict, Any, Optional
from config import *


class DOCXParser:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self.current_chapter = None
        self.current_section = None
        self.extracted_rules = []
        self.extracted_tables = []
        self.extracted_annexures = []
        
    def parse(self) -> Dict[str, Any]:
        """Main parsing method"""
        print(f"\n{'='*80}")
        print(f"Parsing: {self.docx_path.name}")
        print(f"{'='*80}\n")
        
        # Extract all content
        self._extract_paragraphs()
        self._extract_tables()
        
        return {
            'rules': self.extracted_rules,
            'tables': self.extracted_tables,
            'annexures': self.extracted_annexures,
            'stats': {
                'total_rules': len(self.extracted_rules),
                'total_tables': len(self.extracted_tables),
                'total_annexures': len(self.extracted_annexures)
            }
        }
    
    def _extract_paragraphs(self):
        """Extract all paragraphs and identify structure"""
        print("📄 Extracting paragraphs...")
        
        current_rule = None
        rule_text_buffer = []
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for chapter
            chapter_match = self._match_chapter(text)
            if chapter_match:
                self.current_chapter = chapter_match
                print(f"   Found Chapter: {chapter_match['number']} - {chapter_match['title']}")
                continue
            
            # Check for section
            section_match = self._match_section(text)
            if section_match:
                self.current_section = section_match
                print(f"   Found Section: {section_match['number']} - {section_match['title']}")
                continue
            
            # Check for clause/rule
            clause_match = self._match_clause(text)
            if clause_match:
                # Save previous rule if exists
                if current_rule and rule_text_buffer:
                    current_rule['fullText'] = '\n'.join(rule_text_buffer)
                    self.extracted_rules.append(current_rule)
                
                # Start new rule
                current_rule = {
                    'chapter': self.current_chapter['number'] if self.current_chapter else 'Unknown',
                    'chapterTitle': self.current_chapter['title'] if self.current_chapter else '',
                    'section': self.current_section['number'] if self.current_section else '',
                    'sectionTitle': self.current_section['title'] if self.current_section else '',
                    'clause': clause_match['number'],
                    'title': clause_match['title'],
                    'reference': f"{self.current_chapter['number'] if self.current_chapter else 'X'}.{self.current_section['number'] if self.current_section else 'X'}.{clause_match['number']}",
                    'summary': clause_match['title'][:200] if clause_match['title'] else '',
                }
                rule_text_buffer = [text]
                continue
            
            # Check for annexure
            if re.match(r'^ANNEXURE\s+\d+', text, re.IGNORECASE):
                self._extract_annexure(text)
                continue
            
            # Add to current rule buffer
            if current_rule:
                rule_text_buffer.append(text)
        
        # Save last rule
        if current_rule and rule_text_buffer:
            current_rule['fullText'] = '\n'.join(rule_text_buffer)
            self.extracted_rules.append(current_rule)
        
        print(f"   ✅ Extracted {len(self.extracted_rules)} rules\n")
    
    def _extract_tables(self):
        """Extract all tables from document"""
        print("📊 Extracting tables...")
        
        for i, table in enumerate(self.document.tables):
            table_data = self._parse_table(table)
            if table_data:
                self.extracted_tables.append({
                    'table_number': i + 1,
                    'context': self._get_table_context(table),
                    'data': table_data,
                    'chapter': self.current_chapter['number'] if self.current_chapter else 'Unknown',
                    'section': self.current_section['number'] if self.current_section else ''
                })
        
        print(f"   ✅ Extracted {len(self.extracted_tables)} tables\n")
    
    def _parse_table(self, table: Table) -> List[List[str]]:
        """Parse table into 2D array"""
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # Skip empty rows
                data.append(row_data)
        return data if data else None
    
    def _get_table_context(self, table: Table) -> str:
        """Get context around table (previous paragraphs)"""
        # This is a simplified version - you might need to enhance this
        return f"Table in {self.current_chapter['title'] if self.current_chapter else 'Unknown'}"
    
    def _match_chapter(self, text: str) -> Optional[Dict[str, str]]:
        """Match chapter patterns"""
        for pattern in CHAPTER_PATTERNS:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                return {
                    'number': match.group(1),
                    'title': match.group(2).strip() if len(match.groups()) > 1 else ''
                }
        return None
    
    def _match_section(self, text: str) -> Optional[Dict[str, str]]:
        """Match section patterns"""
        for pattern in SECTION_PATTERNS:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                return {
                    'number': match.group(1),
                    'title': match.group(2).strip() if len(match.groups()) > 1 else ''
                }
        return None
    
    def _match_clause(self, text: str) -> Optional[Dict[str, str]]:
        """Match clause/rule patterns"""
        for pattern in CLAUSE_PATTERNS:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                return {
                    'number': match.group(1),
                    'title': match.group(2).strip() if len(match.groups()) > 1 else text
                }
        return None
    
    def _extract_annexure(self, text: str):
        """Extract annexure information"""
        self.extracted_annexures.append({
            'title': text,
            'chapter': self.current_chapter['number'] if self.current_chapter else 'Unknown'
        })


def parse_udcpr() -> Dict[str, Any]:
    """Parse UDCPR document"""
    parser = DOCXParser(UDCPR_DOCX)
    return parser.parse()


def parse_mumbai_dcpr() -> Dict[str, Any]:
    """Parse Mumbai-DCPR document"""
    parser = DOCXParser(MUMBAI_DCPR_DOCX)
    return parser.parse()


if __name__ == '__main__':
    print("\n" + "="*80)
    print("DOCX PARSER TEST")
    print("="*80)
    
    # Test UDCPR
    udcpr_data = parse_udcpr()
    print(f"\nUDCPR Stats:")
    print(f"  Rules: {udcpr_data['stats']['total_rules']}")
    print(f"  Tables: {udcpr_data['stats']['total_tables']}")
    print(f"  Annexures: {udcpr_data['stats']['total_annexures']}")
    
    # Test Mumbai-DCPR
    mumbai_data = parse_mumbai_dcpr()
    print(f"\nMumbai-DCPR Stats:")
    print(f"  Rules: {mumbai_data['stats']['total_rules']}")
    print(f"  Tables: {mumbai_data['stats']['total_tables']}")
    print(f"  Annexures: {mumbai_data['stats']['total_annexures']}")
