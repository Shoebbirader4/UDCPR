"""
Improved DOCX Parser - Extract ALL content as individual rules
Treats each meaningful paragraph/section as a separate rule
"""
from docx import Document
from docx.table import Table
import re
from typing import List, Dict, Any
from config import *


class ImprovedDOCXParser:
    def __init__(self, docx_path: str, source_name: str):
        self.docx_path = docx_path
        self.source_name = source_name
        self.document = Document(docx_path)
        self.extracted_rules = []
        self.rule_counter = 0
        
    def parse(self) -> Dict[str, Any]:
        """Main parsing method - extract everything"""
        print(f"\n{'='*80}")
        print(f"Improved Parsing: {self.docx_path.name}")
        print(f"{'='*80}\n")
        
        # Extract all meaningful content
        self._extract_all_content()
        
        return {
            'rules': self.extracted_rules,
            'stats': {
                'total_rules': len(self.extracted_rules),
                'source': self.source_name
            }
        }
    
    def _extract_all_content(self):
        """Extract all paragraphs and tables as individual rules"""
        print("📄 Extracting all content...")
        
        current_chapter = "General"
        current_section = ""
        paragraph_buffer = []
        
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text or len(text) < 10:
                continue
            
            # Check if this looks like a heading/chapter
            if self._is_heading(para, text):
                # Save previous buffer if exists
                if paragraph_buffer:
                    self._create_rule_from_buffer(paragraph_buffer, current_chapter, current_section)
                    paragraph_buffer = []
                
                # Update chapter/section
                if self._looks_like_chapter(text):
                    current_chapter = text[:100]  # Limit length
                    current_section = ""
                else:
                    current_section = text[:100]
                
                continue
            
            # Check if this is a regulation number pattern
            if self._is_regulation_reference(text):
                # Save previous buffer
                if paragraph_buffer:
                    self._create_rule_from_buffer(paragraph_buffer, current_chapter, current_section)
                    paragraph_buffer = []
                
                # Start new rule with this as title
                paragraph_buffer = [text]
                continue
            
            # Add to buffer
            paragraph_buffer.append(text)
            
            # If buffer gets large, save it as a rule
            if len(paragraph_buffer) >= 5 or len('\n'.join(paragraph_buffer)) > 1000:
                self._create_rule_from_buffer(paragraph_buffer, current_chapter, current_section)
                paragraph_buffer = []
        
        # Save any remaining buffer
        if paragraph_buffer:
            self._create_rule_from_buffer(paragraph_buffer, current_chapter, current_section)
        
        # Also extract tables as separate rules
        self._extract_tables_as_rules()
        
        print(f"   ✅ Extracted {len(self.extracted_rules)} rules\n")
    
    def _is_heading(self, para, text: str) -> bool:
        """Check if paragraph is a heading"""
        # Check style
        if para.style.name.startswith('Heading'):
            return True
        
        # Check if all caps and short
        if text.isupper() and len(text) < 200:
            return True
        
        # Check if bold and short
        if para.runs and len(para.runs) > 0:
            if para.runs[0].bold and len(text) < 200:
                return True
        
        return False
    
    def _looks_like_chapter(self, text: str) -> bool:
        """Check if text looks like a chapter heading"""
        chapter_patterns = [
            r'^CHAPTER\s+[IVX\d]+',
            r'^Chapter\s+[IVX\d]+',
            r'^PART\s+[IVX\d]+',
            r'^Part\s+[IVX\d]+',
        ]
        
        for pattern in chapter_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _is_regulation_reference(self, text: str) -> bool:
        """Check if text starts with a regulation reference"""
        patterns = [
            r'^Regulation\s+No\.\s*\d+',
            r'^Rule\s+\d+',
            r'^\d+\.\d+\s+[A-Z]',
            r'^Section\s+\d+',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _create_rule_from_buffer(self, buffer: List[str], chapter: str, section: str):
        """Create a rule from buffered paragraphs"""
        if not buffer:
            return
        
        full_text = '\n'.join(buffer)
        
        # Skip if too short
        if len(full_text) < 20:
            return
        
        self.rule_counter += 1
        
        # Extract title (first line or first 100 chars)
        title = buffer[0][:200] if buffer else "Untitled"
        
        # Try to extract clause number from title
        clause = self._extract_clause_number(title)
        
        # Create rule
        rule = {
            'chapter': chapter,
            'chapterTitle': chapter,
            'section': section,
            'sectionTitle': section,
            'clause': clause,
            'title': title,
            'reference': f"{self.source_name}-{self.rule_counter}",
            'summary': title[:500],
            'fullText': full_text,
            'source': self.source_name
        }
        
        self.extracted_rules.append(rule)
    
    def _extract_clause_number(self, text: str) -> str:
        """Extract clause/regulation number from text"""
        # Try various patterns
        patterns = [
            r'Regulation\s+No\.\s*(\d+(?:\.\d+)*)',
            r'Rule\s+(\d+(?:\.\d+)*)',
            r'^(\d+\.\d+(?:\.\d+)?)',
            r'Section\s+(\d+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        
        return str(self.rule_counter)
    
    def _extract_tables_as_rules(self):
        """Extract tables as separate rules"""
        print("📊 Extracting tables as rules...")
        
        for i, table in enumerate(self.document.tables):
            table_data = self._parse_table(table)
            if not table_data or len(table_data) < 2:
                continue
            
            # Create a rule from table
            self.rule_counter += 1
            
            # Try to get table title from first row
            title = ' | '.join(table_data[0][:3]) if table_data else f"Table {i+1}"
            
            # Convert table to text
            table_text = self._table_to_text(table_data)
            
            rule = {
                'chapter': 'Tables',
                'chapterTitle': 'Tables and Schedules',
                'section': f'Table {i+1}',
                'sectionTitle': title,
                'clause': f'T{i+1}',
                'title': title[:200],
                'reference': f"{self.source_name}-TABLE-{i+1}",
                'summary': f"Table: {title[:300]}",
                'fullText': table_text,
                'source': self.source_name,
                'isTable': True
            }
            
            self.extracted_rules.append(rule)
        
        print(f"   ✅ Extracted {len(self.document.tables)} tables as rules\n")
    
    def _parse_table(self, table: Table) -> List[List[str]]:
        """Parse table into 2D array"""
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # Skip empty rows
                data.append(row_data)
        return data
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to formatted text"""
        lines = []
        for row in table_data:
            lines.append(' | '.join(row))
        return '\n'.join(lines)


def parse_udcpr_improved() -> Dict[str, Any]:
    """Parse UDCPR with improved method"""
    parser = ImprovedDOCXParser(UDCPR_DOCX, 'UDCPR')
    return parser.parse()


def parse_mumbai_dcpr_improved() -> Dict[str, Any]:
    """Parse Mumbai-DCPR with improved method"""
    parser = ImprovedDOCXParser(MUMBAI_DCPR_DOCX, 'Mumbai-DCPR')
    return parser.parse()


if __name__ == '__main__':
    print("\n" + "="*80)
    print("IMPROVED DOCX PARSER TEST")
    print("="*80)
    
    # Test UDCPR
    udcpr_data = parse_udcpr_improved()
    print(f"\nUDCPR Stats:")
    print(f"  Rules: {udcpr_data['stats']['total_rules']}")
    
    # Test Mumbai-DCPR
    mumbai_data = parse_mumbai_dcpr_improved()
    print(f"\nMumbai-DCPR Stats:")
    print(f"  Rules: {mumbai_data['stats']['total_rules']}")
    
    print(f"\nTotal: {udcpr_data['stats']['total_rules'] + mumbai_data['stats']['total_rules']} rules")
